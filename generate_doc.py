"""
Generate comprehensive OpenEnv hackathon preparation guide as DOCX.
Run with: py generate_doc.py
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
section = doc.sections[0]
section.top_margin    = Cm(2.0)
section.bottom_margin = Cm(2.0)
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2.5)

# ── Helper colours ────────────────────────────────────────────────────────────
DARK_BLUE   = RGBColor(0x1A, 0x3A, 0x5C)
ACCENT_BLUE = RGBColor(0x1E, 0x6E, 0xC8)
ACCENT_ORG  = RGBColor(0xE0, 0x6C, 0x00)
WHITE       = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_GRAY  = RGBColor(0xF2, 0xF2, 0xF2)

def set_cell_bg(cell, hex_color: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)

def heading(text, level=1, color=DARK_BLUE):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = color
        run.font.bold = True
        if level == 1:
            run.font.size = Pt(18)
        elif level == 2:
            run.font.size = Pt(14)
        else:
            run.font.size = Pt(12)
    p.paragraph_format.space_before = Pt(14 if level==1 else 10)
    p.paragraph_format.space_after  = Pt(6)
    return p

def body(text, bold=False, italic=False, color=None, size=11, indent=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(4)
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    run = p.add_run(text)
    run.font.size   = Pt(size)
    run.font.bold   = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color
    return p

def bullet(text, level=0, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent   = Cm(0.5 + level * 0.5)
    p.paragraph_format.space_before  = Pt(1)
    p.paragraph_format.space_after   = Pt(2)
    if bold_prefix:
        r = p.add_run(bold_prefix + " ")
        r.bold = True
        r.font.size = Pt(11)
    r2 = p.add_run(text)
    r2.font.size = Pt(11)
    return p

def numbered(text, level=0):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.left_indent  = Cm(0.5 + level * 0.5)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.font.size = Pt(11)
    return p

def code_block(text):
    p  = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(1.0)
    p.paragraph_format.right_indent = Cm(1.0)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1E, 0x40, 0x80)
    return p

def divider():
    doc.add_paragraph("─" * 90)

def callout(label, text, bg="E8F0FE"):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.rows[0].cells[0]
    set_cell_bg(cell, bg)
    p = cell.paragraphs[0]
    r1 = p.add_run(label + "  ")
    r1.bold = True
    r1.font.size = Pt(11)
    r1.font.color.rgb = DARK_BLUE
    r2 = p.add_run(text)
    r2.font.size = Pt(11)
    doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════════════════════
#  COVER PAGE
# ═══════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("OPEN ENVIRONMENT HACKATHON")
r.font.size  = Pt(28)
r.font.bold  = True
r.font.color.rgb = DARK_BLUE

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Complete Preparation Guide — Round 2")
r.font.size  = Pt(18)
r.font.color.rgb = ACCENT_BLUE

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("ATC Runway Optimisation  ·  OpenEnv  ·  Reinforcement Learning")
r.font.size  = Pt(13)
r.font.italic = True
r.font.color.rgb = ACCENT_ORG

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Ganeev Singh  |  April 2026")
r.font.size  = Pt(11)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  TABLE OF CONTENTS  (manual)
# ═══════════════════════════════════════════════════════════════════════════════
heading("Table of Contents", level=1)
toc_items = [
    ("1",  "What Is an Open Environment (OpenEnv)?"),
    ("2",  "Reinforcement Learning Core Concepts"),
    ("3",  "Your Round-1 Project — Architecture Deep Dive"),
    ("4",  "File-by-File Knowledge Map"),
    ("5",  "The 3-Layer Grading System — Master This"),
    ("6",  "Reward Shaping & Potential-Based Rewards"),
    ("7",  "Wake Turbulence & Aviation Domain Knowledge"),
    ("8",  "What Stays Fixed vs. What Changes in Round 2"),
    ("9",  "Round 2 Innovation Playbook"),
    ("10", "Explaining to Judges — Cheat Sheet"),
    ("11", "Key Formulas & Numbers to Memorise"),
    ("12", "Coding Patterns You Must Know"),
]
for num, title in toc_items:
    p = doc.add_paragraph()
    r = p.add_run(f"  {num}.  {title}")
    r.font.size = Pt(11)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 1 — WHAT IS AN OPEN ENVIRONMENT
# ═══════════════════════════════════════════════════════════════════════════════
heading("1. What Is an Open Environment (OpenEnv)?", level=1)

body("Definition", bold=True, color=DARK_BLUE, size=12)
body(
    "An Open Environment (OpenEnv) is a standardised, HTTP-based reinforcement-learning "
    "environment. Any agent — an LLM, a search algorithm, a human — can interact with it "
    "through three universal endpoints: /reset, /step, and /state. "
    "OpenEnv is to RL what REST is to web services: a shared contract everyone agrees on."
)

body("The Three Core API Endpoints", bold=True, color=DARK_BLUE, size=12)
table = doc.add_table(rows=4, cols=3)
table.style = "Table Grid"
hdr_cells = table.rows[0].cells
for i, h in enumerate(["Endpoint", "Purpose", "Returns"]):
    hdr_cells[i].text = h
    set_cell_bg(hdr_cells[i], "1A3A5C")
    for r in hdr_cells[i].paragraphs[0].runs:
        r.font.color.rgb = WHITE
        r.font.bold = True

data = [
    ("POST /reset", "Start a new episode; choose task / difficulty / seed", "Observation (initial state briefing)"),
    ("POST /step",  "Submit an action; advance the environment one step",   "Observation (new state + reward + done flag)"),
    ("GET  /state", "Inspect current internal state (read-only)",           "State object (full history, best plan)"),
]
for row_idx, (ep, pur, ret) in enumerate(data, 1):
    cells = table.rows[row_idx].cells
    cells[0].text = ep
    cells[1].text = pur
    cells[2].text = ret
doc.add_paragraph()

body("OpenEnv vs. OpenAI Gym — Key Difference", bold=True, color=DARK_BLUE, size=12)
body(
    "OpenAI Gym environments live inside your process (Python objects). "
    "OpenEnv environments are network services — agents talk HTTP, actions are JSON. "
    "This means any language, any LLM, and multiple concurrent agents can share one environment."
)

callout("JUDGE TIP:", "When a judge asks 'why OpenEnv?'  Answer: standardised benchmark contract, "
        "multi-agent capable, language-agnostic, deployable to cloud/HF Space.", "E8F4FD")

body("The openenv.yaml — The Environment's Passport", bold=True, color=DARK_BLUE, size=12)
bullet("Declares spec_version, name, runtime, action_space, observation_space")
bullet("Lists tasks with programmatic_grader flags and max_steps")
bullet("Used by OpenEnv core to auto-wire FastAPI routes and validation")
code_block(
    "spec_version: 1\n"
    "name: atc_env\n"
    "runtime: fastapi\n"
    "app: server.app:app\n"
    "action_space:  models:ATCOptimizationAction\n"
    "observation_space: models:ATCOptimizationObservation\n"
    "tasks:\n"
    "  - id: delhi_monsoon_recovery_easy\n"
    "    programmatic_grader: true\n"
    "    max_steps: 8"
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 2 — RL CORE CONCEPTS
# ═══════════════════════════════════════════════════════════════════════════════
heading("2. Reinforcement Learning Core Concepts", level=1)

body("The RL Loop (must know cold)", bold=True, color=DARK_BLUE, size=12)
code_block(
    "Agent  ──action──►  Environment\n"
    "       ◄──(obs, reward, done)──"
)
body("At every timestep t:")
bullet("Agent receives observation  o_t  (what it can see)")
bullet("Agent chooses action  a_t  (its decision)")
bullet("Environment transitions to new state  s_{t+1}")
bullet("Agent receives reward  r_t  (scalar feedback)")
bullet("Episode ends when done=True")

body("Key RL Vocabulary", bold=True, color=DARK_BLUE, size=12)
table = doc.add_table(rows=11, cols=2)
table.style = "Table Grid"
for i, h in enumerate(["Term", "Plain-English Meaning"]):
    table.rows[0].cells[i].text = h
    set_cell_bg(table.rows[0].cells[i], "1A3A5C")
    for r in table.rows[0].cells[i].paragraphs[0].runs:
        r.font.color.rgb = WHITE; r.font.bold = True

terms = [
    ("State (s)",         "Complete internal representation of the environment at time t"),
    ("Observation (o)",   "What the agent actually sees (may be partial view of full state)"),
    ("Action (a)",        "Decision made by the agent — in your project: a list of SlotAssignments"),
    ("Reward (r)",        "Scalar signal: how good was the last action? (potential-based in your env)"),
    ("Policy (π)",        "Agent's strategy: maps observation → action"),
    ("Episode",           "One complete run from reset to done"),
    ("Trajectory",        "Sequence of (s, a, r) tuples across an episode"),
    ("Return (G)",        "Sum of future rewards, usually discounted: Σ γ^t · r_t"),
    ("Value function V(s)","Expected return from state s following policy π"),
    ("Done / Terminal",   "Boolean signal that the episode has ended"),
]
for row_idx, (term, meaning) in enumerate(terms, 1):
    table.rows[row_idx].cells[0].text = term
    table.rows[row_idx].cells[1].text = meaning
doc.add_paragraph()

body("Types of Environments You Should Know", bold=True, color=DARK_BLUE, size=12)
bullet("Discrete Action Space — finite set of moves (chess, Atari)")
bullet("Continuous Action Space — real-valued actions (robotics, control)")
bullet("Combinatorial Action Space — YOUR environment: assign N flights to runways/slots")
bullet("Sparse Reward — reward only at episode end (hard for learning)")
bullet("Dense Reward — reward every step (your env uses potential-based shaping for this)")
bullet("Deterministic — same action always → same result (your scoring engine is deterministic)")
bullet("Stochastic — randomness in transitions (weather, sensor noise)")

body("Why Dense Reward Matters for LLM Agents", bold=True, color=DARK_BLUE, size=12)
body(
    "LLM-based agents cannot do gradient descent, so they rely on feedback in the observation "
    "to improve their next action. Dense per-step rewards let the agent see 'am I improving?' "
    "after every iteration. Your environment does this via potential-based shaping."
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 3 — PROJECT ARCHITECTURE
# ═══════════════════════════════════════════════════════════════════════════════
heading("3. Your Round-1 Project — Architecture Deep Dive", level=1)

body("One-Line Pitch", bold=True, color=DARK_BLUE, size=12)
body(
    "A production-ready OpenEnv benchmark where an AI agent must re-sequence disrupted flights "
    "onto runways under real-world constraints (wake turbulence safety, emergency priority, "
    "airline fairness, fuel burn) — scored by a 3-layer gated composite grader."
)

body("Data Flow — Step by Step", bold=True, color=DARK_BLUE, size=12)
numbered("Agent sends POST /reset → environment initialises scenario → returns observation with briefing")
numbered("Agent reads observation: flights list, runways, metrics (all zeros), recommendations")
numbered("Agent crafts ATCOptimizationAction: list of SlotAssignments + rationale + commit=False")
numbered("Agent sends POST /step with action")
numbered("engine.py simulates_plan() → calculates 8 metric dimensions")
numbered("graders.py GatedCompositeGrader() → official score (0,1)")
numbered("atc_environment.py computes potential-based reward = current_score − previous_score")
numbered("Observation returned: new metrics, diagnostics, recommendations, reward, done")
numbered("Agent refines plan based on feedback; repeats until commit=True or max_steps")

body("The Six Constraint Categories", bold=True, color=DARK_BLUE, size=12)
table = doc.add_table(rows=7, cols=3)
table.style = "Table Grid"
for i, h in enumerate(["Constraint", "Source of Truth", "Violation Penalty"]):
    table.rows[0].cells[i].text = h
    set_cell_bg(table.rows[0].cells[i], "1A3A5C")
    for r in table.rows[0].cells[i].paragraphs[0].runs:
        r.font.color.rgb = WHITE; r.font.bold = True

constraints = [
    ("Wake Turbulence Separation", "constants.py SEPARATION_BY_WAKE", "conflict_free_ratio drops; safety gate ceiling applied"),
    ("Time Window [earliest, latest]", "tasks.py FlightRecord", "Flight excluded from completeness score"),
    ("Runway Capability Match",       "tasks.py RunwaySpec.allowed_ops", "Invalid assignment rejected"),
    ("Priority Delay Tolerance",      "engine.py PRIORITY_TOLERANCES",  "priority_handling score penalised"),
    ("Airline Fairness",              "engine.py std-dev formula",       "fairness score penalised"),
    ("Fuel Burn Budget",              "tasks.py TaskDefinition.budget",  "fuel_efficiency score penalised"),
]
for row_idx, (c, s, p) in enumerate(constraints, 1):
    table.rows[row_idx].cells[0].text = c
    table.rows[row_idx].cells[1].text = s
    table.rows[row_idx].cells[2].text = p
doc.add_paragraph()

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 4 — FILE-BY-FILE KNOWLEDGE MAP
# ═══════════════════════════════════════════════════════════════════════════════
heading("4. File-by-File Knowledge Map", level=1)
body("What every file does and why it exists:", bold=True, color=DARK_BLUE, size=12)

files = [
    ("models.py",
     "Pydantic typed contracts — the shared vocabulary of the entire system.",
     ["ATCOptimizationAction: what the agent sends (proposals + rationale + commit)",
      "ATCOptimizationObservation: what the agent receives (metrics + diagnostics + reward)",
      "ATCOptimizationState: internal snapshot exposed via /state",
      "SlotAssignment: one flight → (runway, minute, hold)",
      "TaskGrade: official score + sub_scores + rationale"]),

    ("tasks.py",
     "Scenario catalog — the problem instances the agent must solve.",
     ["4 tasks: easy (10 flights) → medium (13) → hard (17) → medium-hard (7 + 1 runway)",
      "render_task_briefing(): text prompt given to LLM agent",
      "task_catalog(): dict for /reset endpoint to look up tasks by ID"]),

    ("constants.py",
     "Single source of truth for all magic numbers.",
     ["SEPARATION_BY_WAKE: 3×3 matrix (Heavy/Medium/Light) → minimum gap in minutes",
      "SCORE_WEIGHTS: completeness(24%), conflict(24%), priority(18%), delay(16%), fairness(10%), fuel(8%)",
      "Inference defaults: MAX_STEPS=4, MAX_TOKENS=1400, TEMP=0"]),

    ("engine.py",
     "Deterministic simulator — converts a plan into a score.",
     ["simulate_plan(task, proposals) → SimulationOutcome",
      "Validates: duplicates, unknown flights, runway compatibility, time windows",
      "Detects separation conflicts using wake matrix",
      "Calculates all 6 metric dimensions + normalized_score",
      "Returns diagnostics list (up to 12) and recommendations list (up to 6)"]),

    ("graders.py",
     "3-layer gated composite grader — the official benchmark scorer.",
     ["Layer 1 SafetyGateEvaluator: hard ceiling on score if conflicts/missing flights",
      "Layer 2 PriorityRubricGrader: 30% of score — emergency/medical/connection handling",
      "Layer 3 EfficiencyRubricGrader: 70% of score — delay/fuel/fairness/connection impact",
      "GatedCompositeGrader: min(gate_ceiling, 0.3×priority + 0.7×efficiency)",
      "LLMSupervisorGrader: auxiliary only, calls OpenAI-compatible API for narrative feedback"]),

    ("planner.py",
     "Deterministic baseline planner — fallback when LLM unavailable.",
     ["build_heuristic_plan(): greedy by priority rank → first conflict-free schedule",
      "build_refined_plan(): iterative 2-pass improvement using simulate_plan feedback",
      "Used by inference.py as seed plan passed to LLM"]),

    ("server/atc_environment.py",
     "The OpenEnv Environment class — implements the RL loop.",
     ["ATCOptimizationEnvironment(Environment[Action, Observation, State])",
      "reset(): initialise episode, return briefing observation",
      "step(): simulate → grade → compute potential-based reward → return observation",
      "state property: serialisable snapshot with full history",
      "SUPPORTS_CONCURRENT_SESSIONS = True"]),

    ("server/app.py",
     "FastAPI entrypoint — wires everything to HTTP.",
     ["create_app(ATCOptimizationEnvironment, max_environments=8)",
      "GET /: browser UI with task strips and model options",
      "POST /ui/run-inference: browser-triggered inference"]),

    ("inference.py",
     "Baseline inference script — shows how an LLM agent runs.",
     ["Launches uvicorn server subprocess if no external URL",
      "For each task: reset → loop(get_model_action → step) → log [START/STEP/END]",
      "get_model_action(): builds seed plan → queries LLM → falls back to deterministic",
      "Writes final avg score to stdout in strict format"]),

    ("client.py",
     "HTTP client wrapper for the environment.",
     ["ATCOptimizationEnv(EnvClient[Action, Observation, State])",
      "_step_payload(): serialises action to JSON",
      "_parse_result(): deserialises observation from JSON"]),

    ("openenv.yaml",
     "Environment passport — OpenEnv metadata spec.",
     ["Declares action/observation spaces, task list, grader type, max_steps"]),
]

for fname, purpose, details in files:
    heading(fname, level=2, color=ACCENT_BLUE)
    body(purpose, italic=True)
    for d in details:
        bullet(d)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 5 — THE 3-LAYER GRADING SYSTEM
# ═══════════════════════════════════════════════════════════════════════════════
heading("5. The 3-Layer Grading System — Master This", level=1)

body(
    "This is the most important design decision in your project. Judges WILL ask about it. "
    "Know every layer cold.", italic=True, color=ACCENT_ORG
)

body("Why Three Layers?", bold=True, color=DARK_BLUE, size=12)
body(
    "Real-world aviation has hard safety requirements that cannot be traded off against efficiency. "
    "A plane cannot be 'a little bit in conflict' — it either is or it isn't. "
    "The 3-layer architecture enforces this: safety violations cap the maximum achievable score, "
    "regardless of how efficient or fair the schedule is."
)

heading("Layer 1 — Safety Gate (Hard Ceiling)", level=2, color=ACCENT_ORG)
body("Applied FIRST. Violations reduce the maximum score the agent can achieve.")
table = doc.add_table(rows=4, cols=3)
table.style = "Table Grid"
for i, h in enumerate(["Violation", "Ceiling Formula", "Example"]):
    table.rows[0].cells[i].text = h
    set_cell_bg(table.rows[0].cells[i], "E06C00")
    for r in table.rows[0].cells[i].paragraphs[0].runs:
        r.font.color.rgb = WHITE; r.font.bold = True
safety_data = [
    ("Separation conflict",    "max(0.10, 0.40 − 0.05 × (n_conflicts − 1))", "2 conflicts → ceiling = 0.35"),
    ("Missing assignments",    "max(0.20, 0.50 − 0.04 × missing_count)",      "3 missing → ceiling = 0.38"),
    ("Emergency delay > 5 min","Fixed ceiling = 0.35",                         "1 delayed emergency → cap 0.35"),
]
for row_idx, (v, f, e) in enumerate(safety_data, 1):
    table.rows[row_idx].cells[0].text = v
    table.rows[row_idx].cells[1].text = f
    table.rows[row_idx].cells[2].text = e
doc.add_paragraph()

heading("Layer 2 — Priority Rubric (30% of Final Score)", level=2, color=ACCENT_BLUE)
body("Measures how well the agent handles priority flights.")
code_block(
    "priority_score = 0.50 × emergency_score\n"
    "               + 0.30 × medical_score\n"
    "               + 0.20 × connection_score\n\n"
    "emergency_score  = 1.0  if no delays > 5 min,  else 0.0  (binary!)\n"
    "medical_score    = 1.0 − (violated_count / total_medical)\n"
    "connection_score = from engine.py connection_impact_score"
)

heading("Layer 3 — Efficiency Rubric (70% of Final Score)", level=2, color=ACCENT_BLUE)
body("Measures operational efficiency.")
code_block(
    "efficiency_score = 0.35 × delay_efficiency\n"
    "                 + 0.25 × fuel_efficiency\n"
    "                 + 0.20 × fairness_score\n"
    "                 + 0.20 × connection_impact_score"
)

heading("Composite Formula", level=2, color=DARK_BLUE)
code_block(
    "final_score = min(safety_gate_ceiling,\n"
    "                  0.30 × priority_score  +  0.70 × efficiency_score)\n\n"
    "# Always clipped to (0.01, 0.99) — never exactly 0 or 1"
)
callout("KEY INSIGHT:", "The min() means safety can ONLY hurt, never help. A perfect efficiency "
        "score is worthless if you have a separation conflict. This mirrors real aviation: "
        "safety is non-negotiable.", "FFF3CD")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 6 — REWARD SHAPING
# ═══════════════════════════════════════════════════════════════════════════════
heading("6. Reward Shaping & Potential-Based Rewards", level=1)

body("The Problem With Sparse Rewards", bold=True, color=DARK_BLUE, size=12)
body(
    "If an agent only receives reward at the end of an episode (sparse), it has no signal "
    "during the episode to know if it is improving. For an LLM that cannot do gradient descent, "
    "this is especially bad — it needs per-step feedback to refine its plan."
)

body("Potential-Based Reward Shaping (Ng et al. 1999)", bold=True, color=DARK_BLUE, size=12)
body(
    "This is the theoretically safe way to add dense rewards without changing the optimal policy. "
    "The idea: define a potential function Φ(s) over states. "
    "The shaped reward is:"
)
code_block(
    "r_shaped(s, a, s') = r_original(s, a, s')  +  γ·Φ(s') − Φ(s)\n\n"
    "In your environment (γ=1, terminal reward only):\n"
    "  Φ(s) = composite_score(s)\n"
    "  r_t   = score_t  −  score_{t-1}   (delta between consecutive steps)\n\n"
    "Property: sum of shaped rewards over episode = final_score − initial_score (0)\n"
    "          = final_score   (same as sparse reward — policy invariant!)"
)
callout("WHY THIS MATTERS:", "Judges from ML backgrounds will be impressed you know potential-based "
        "shaping. It proves your reward design is theoretically grounded, not ad-hoc.", "E8F8E8")

body("Your Implementation (atc_environment.py)", bold=True, color=DARK_BLUE, size=12)
code_block(
    "# In step():\n"
    "current_score   = grader_result.score          # from GatedCompositeGrader\n"
    "reward          = current_score - self._previous_score\n"
    "reward          = max(-1.0, min(1.0, reward))  # clamp to [-1, 1]\n"
    "self._previous_score = current_score\n\n"
    "# Done conditions:\n"
    "done = action.commit or step_count >= max_steps or current_score >= 0.98"
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 7 — AVIATION DOMAIN KNOWLEDGE
# ═══════════════════════════════════════════════════════════════════════════════
heading("7. Wake Turbulence & Aviation Domain Knowledge", level=1)

body("Wake Turbulence — Why It Matters", bold=True, color=DARK_BLUE, size=12)
body(
    "Every aircraft generates a vortex wake behind it. Lighter aircraft following heavier ones "
    "can be violently rolled. ICAO (International Civil Aviation Organization) mandates minimum "
    "separation times between aircraft of different weight classes."
)

body("Wake Separation Matrix (minutes)", bold=True, color=DARK_BLUE, size=12)
table = doc.add_table(rows=4, cols=4)
table.style = "Table Grid"
headers = ["Leader → Follower", "Follower: HEAVY", "Follower: MEDIUM", "Follower: LIGHT"]
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h
    set_cell_bg(table.rows[0].cells[i], "1A3A5C")
    for r in table.rows[0].cells[i].paragraphs[0].runs:
        r.font.color.rgb = WHITE; r.font.bold = True
wake_data = [
    ("Leader: HEAVY",  "4 min", "5 min", "6 min"),
    ("Leader: MEDIUM", "3 min", "3 min", "4 min"),
    ("Leader: LIGHT",  "3 min", "3 min", "3 min"),
]
for row_idx, row in enumerate(wake_data, 1):
    for col_idx, val in enumerate(row):
        table.rows[row_idx].cells[col_idx].text = val
doc.add_paragraph()

body("Key Asymmetry", bold=True, color=DARK_BLUE, size=12)
body(
    "LIGHT after HEAVY requires 6 minutes — the longest gap. "
    "HEAVY after LIGHT only needs 3 minutes. "
    "Optimal sequencing exploits this: put LIGHT aircraft BEFORE heavy ones to minimise total wait time."
)
callout("EXAM QUESTION ANSWER:", "Why does your Hyderabad task specifically reward optimal ordering? "
        "Because 3 wake classes are present and the LIGHT→HEAVY ordering saves ~6 minutes of total delay "
        "compared to HEAVY→LIGHT ordering.", "FFF3CD")

body("Priority Classes (hard constraints)", bold=True, color=DARK_BLUE, size=12)
table = doc.add_table(rows=5, cols=3)
table.style = "Table Grid"
for i, h in enumerate(["Priority Class", "Max Delay Tolerance", "Real-World Meaning"]):
    table.rows[0].cells[i].text = h
    set_cell_bg(table.rows[0].cells[i], "1A3A5C")
    for r in table.rows[0].cells[i].paragraphs[0].runs:
        r.font.color.rgb = WHITE; r.font.bold = True
priority_data = [
    ("EMERGENCY", "5 minutes",  "Distress, fuel critical, medical emergency in-flight"),
    ("MEDICAL",   "10 minutes", "Organ transport, patient on board"),
    ("CONNECTION","20 minutes", "Passenger bank with tight onward flight"),
    ("NORMAL",    "35 minutes", "Standard commercial operation"),
]
for row_idx, (pc, tol, meaning) in enumerate(priority_data, 1):
    table.rows[row_idx].cells[0].text = pc
    table.rows[row_idx].cells[1].text = tol
    table.rows[row_idx].cells[2].text = meaning
doc.add_paragraph()

body("Other Domain Terms", bold=True, color=DARK_BLUE, size=12)
bullet("IRROPS (Irregular Operations): Unplanned disruption — weather, bird-strike, technical")
bullet("Hub Bank: Multiple flights arriving/departing within a short window for connections")
bullet("IATA Worldwide Slot Guidelines (WSG): International framework for slot fairness")
bullet("VFR (Visual Flight Rules): Small aircraft flying by visual reference, no instrument clearance")
bullet("ATFM (Air Traffic Flow Management): European system for pre-departure slot coordination")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 8 — WHAT STAYS FIXED vs. WHAT CHANGES
# ═══════════════════════════════════════════════════════════════════════════════
heading("8. What Stays Fixed vs. What Changes in Round 2", level=1)

body(
    "In Round 2 you will receive a new problem statement and must build a new OpenEnv on the spot. "
    "Understanding which parts of your architecture are reusable is critical.", italic=True
)

body("FIXED — Reuse Without Change", bold=True, color=ACCENT_BLUE, size=13)
table = doc.add_table(rows=9, cols=2)
table.style = "Table Grid"
for i, h in enumerate(["Component", "Why It Stays the Same"]):
    table.rows[0].cells[i].text = h
    set_cell_bg(table.rows[0].cells[i], "1E6EC8")
    for r in table.rows[0].cells[i].paragraphs[0].runs:
        r.font.color.rgb = WHITE; r.font.bold = True
fixed = [
    ("OpenEnv HTTP contract (/reset, /step, /state)", "The standard is fixed — always use these endpoints"),
    ("FastAPI + create_app() pattern",                "OpenEnv core handles routing — just pass your Environment class"),
    ("openenv.yaml structure",                        "Metadata format is standardised — just change values"),
    ("Pydantic models pattern",                       "Always define Action, Observation, State as Pydantic models"),
    ("Environment[Action, Observation, State] base",  "Always inherit from this; implement reset() and step()"),
    ("Potential-based reward shaping formula",        "Works for any scoring function — just swap Φ()"),
    ("Structured logging [START/STEP/END]",           "Required by OpenEnv eval harness — keep format"),
    ("Dockerfile + uvicorn deployment pattern",       "Same deployment target (HF Space) — same container setup"),
]
for row_idx, (c, r) in enumerate(fixed, 1):
    table.rows[row_idx].cells[0].text = c
    table.rows[row_idx].cells[1].text = r
doc.add_paragraph()

body("CHANGES — Replace With New Problem Domain", bold=True, color=ACCENT_ORG, size=13)
table = doc.add_table(rows=9, cols=2)
table.style = "Table Grid"
for i, h in enumerate(["Component", "How to Adapt It"]):
    table.rows[0].cells[i].text = h
    set_cell_bg(table.rows[0].cells[i], "E06C00")
    for r in table.rows[0].cells[i].paragraphs[0].runs:
        r.font.color.rgb = WHITE; r.font.bold = True
changes = [
    ("models.py",           "Redefine Action, Observation, State for your new domain"),
    ("tasks.py",            "Replace with new scenario catalog matching new problem"),
    ("constants.py",        "Replace domain-specific numbers (separation times → your domain's constraints)"),
    ("engine.py",           "Rewrite simulate_plan() for new domain physics/logic"),
    ("graders.py",          "Redesign 3-layer grader for new constraint hierarchy"),
    ("planner.py",          "Write new heuristic baseline for your domain"),
    ("openenv.yaml tasks",  "List new task IDs, max_steps, difficulty levels"),
    ("Task briefing text",  "Write new prompt given to LLM agent describing new scenario"),
]
for row_idx, (c, r) in enumerate(changes, 1):
    table.rows[row_idx].cells[0].text = c
    table.rows[row_idx].cells[1].text = r
doc.add_paragraph()

callout("ROUND 2 STRATEGY:", "Copy your entire repo as a template. Strip out the ATC domain "
        "(tasks.py, engine.py, graders.py, constants.py, models.py). "
        "Keep all the infrastructure. Rewrite only the domain logic.", "E8F8E8")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 9 — INNOVATION PLAYBOOK
# ═══════════════════════════════════════════════════════════════════════════════
heading("9. Round 2 Innovation Playbook", level=1)

body(
    "These are specific innovations you can apply to any new OpenEnv problem. "
    "Each one differentiates you from teams that build a bare-minimum environment.",
    italic=True
)

heading("9.1 — 3-Layer Gated Grader (your signature)", level=2, color=DARK_BLUE)
body(
    "Whatever the new domain, identify the constraint hierarchy: "
    "what are the hard safety constraints (gate), what is critical (30%), what is efficiency (70%)? "
    "Apply the same min(gate, weighted_combo) formula."
)
bullet("Domain: Hospital scheduling → Safety gate: no patient left without care; Priority: critical patients; Efficiency: wait times")
bullet("Domain: Supply chain → Safety gate: mandatory orders not missed; Priority: perishable goods; Efficiency: cost/time")
bullet("Domain: Power grid → Safety gate: no blackout; Priority: hospitals/emergency services; Efficiency: cost")

heading("9.2 — Potential-Based Dense Rewards", level=2, color=DARK_BLUE)
body(
    "Always use reward = score_t − score_{t-1}. "
    "This works for any scoring function and is theoretically justified. "
    "Mention Ng et al. 1999 to impress judges."
)

heading("9.3 — Seed Plan + LLM Refinement Loop", level=2, color=DARK_BLUE)
body(
    "Always build a deterministic heuristic baseline first (like planner.py). "
    "Pass the seed plan as JSON context to the LLM so it refines rather than generates from scratch. "
    "This pattern works for any combinatorial optimisation problem."
)
code_block(
    "# Generic LLM refinement loop\n"
    "seed_plan = heuristic_planner(task)\n"
    "for step in range(MAX_STEPS):\n"
    "    prompt = briefing + current_metrics + seed_plan_json\n"
    "    action = llm.complete(prompt)  or  seed_plan\n"
    "    obs    = env.step(action)\n"
    "    if obs.reward > 0:\n"
    "        seed_plan = action  # keep improvements\n"
    "    if obs.done:\n"
    "        break"
)

heading("9.4 — Diagnostics-as-Observations", level=2, color=DARK_BLUE)
body(
    "Your environment returns structured diagnostics like 'Flight AIC845 assigned to runway incompatible "
    "with operation type ARRIVAL'. These guide the LLM's next action. "
    "For any new domain, invest in generating rich diagnostic messages — they turn the environment "
    "into a self-teaching system."
)

heading("9.5 — Difficulty Progression (easy → hard)", level=2, color=DARK_BLUE)
body(
    "Always create at least 3 tasks of increasing difficulty. This allows evaluators to measure "
    "agent capability at multiple levels and makes your benchmark more useful."
)
bullet("Easy: minimal constraints, 1 safety concern, generous time windows")
bullet("Medium: 2-3 competing constraints, tight time windows, fairness requirement")
bullet("Hard: all constraints active simultaneously, minimal steps, emergency conditions")

heading("9.6 — Concurrent Sessions (SUPPORTS_CONCURRENT_SESSIONS = True)", level=2, color=DARK_BLUE)
body(
    "Always set this flag. It allows multiple agents to evaluate simultaneously. "
    "Technically requires: no global mutable state; all state scoped to episode_id."
)

heading("9.7 — Deterministic Scoring", level=2, color=DARK_BLUE)
body(
    "Your engine.py is fully deterministic — same input always produces same score. "
    "This is critical for reproducible benchmarking. "
    "For any new domain, ensure your simulation has no random elements unless you seed and document them."
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 10 — JUDGE CHEAT SHEET
# ═══════════════════════════════════════════════════════════════════════════════
heading("10. Explaining to Judges — Cheat Sheet", level=1)

questions = [
    ("What is an Open Environment?",
     "A standardised HTTP-based RL environment. Agents interact via /reset, /step, /state. "
     "Any language, any model, any agent type can use the same API. "
     "Like OpenAI Gym but as a network service."),

    ("Why did you choose air traffic control?",
     "It's a real, high-stakes combinatorial optimisation problem. "
     "Hard constraints (safety), soft constraints (efficiency), and priority hierarchies "
     "make it ideal for demonstrating a nuanced grading architecture. "
     "Every constraint has a real-world aviation counterpart."),

    ("Explain your grading system.",
     "3 layers: Safety gate first — violations cap the maximum score regardless of efficiency. "
     "Then priority rubric (30%) for emergency/medical handling. "
     "Then efficiency rubric (70%) for delay/fuel/fairness. "
     "Final = min(gate_ceiling, 0.3×priority + 0.7×efficiency). "
     "This mirrors aviation: safety is non-negotiable."),

    ("How does your reward work?",
     "Potential-based shaping (Ng et al. 1999). "
     "reward_t = composite_score_t − composite_score_{t-1}. "
     "Dense per-step signal, policy-gradient safe, "
     "sum over episode equals final score (same as sparse reward)."),

    ("How does the LLM agent work?",
     "1. Build deterministic seed plan (heuristic). "
     "2. Feed seed plan + task briefing + current metrics to LLM. "
     "3. LLM returns JSON: list of SlotAssignments + rationale + commit flag. "
     "4. Environment scores it, returns diagnostics and new metrics. "
     "5. Repeat up to MAX_STEPS times. Falls back to deterministic plan if LLM fails."),

    ("What is your baseline score?",
     "Heuristic baseline: avg 0.9134 across 4 tasks. "
     "Random baseline: avg 0.165. "
     "Improvement factor: 5.5× over random. "
     "Runtime: 11.69 seconds on 2 vCPU / 8 GB."),

    ("How does wake turbulence affect your scoring?",
     "Heavy→Light requires 6 minutes separation (most restrictive). "
     "Light→Heavy only 3 minutes. "
     "Optimal sequencing puts LIGHT aircraft BEFORE heavy ones to minimise total delay. "
     "Violations trigger the safety gate ceiling — score capped regardless of other metrics."),

    ("What makes your environment production-ready?",
     "Deterministic scoring (reproducible). "
     "Strict score bounds (0.01, 0.99). "
     "Concurrent session support. "
     "Structured logging for eval harness. "
     "Docker + HF Space deployment. "
     "Full pytest suite. "
     "openenv.yaml metadata spec."),
]

for q, a in questions:
    body(f"Q: {q}", bold=True, color=DARK_BLUE, size=12)
    body(f"A: {a}", color=RGBColor(0x20, 0x20, 0x20))
    doc.add_paragraph()

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 11 — KEY FORMULAS
# ═══════════════════════════════════════════════════════════════════════════════
heading("11. Key Formulas & Numbers to Memorise", level=1)

body("Score Weights (memorise these)", bold=True, color=DARK_BLUE, size=12)
code_block(
    "completeness:        0.24  (24%)\n"
    "conflict_free_ratio: 0.24  (24%)\n"
    "priority_handling:   0.18  (18%)\n"
    "delay_efficiency:    0.16  (16%)\n"
    "fairness:            0.10  (10%)\n"
    "fuel_efficiency:     0.08  ( 8%)\n"
    "────────────────────────────────\n"
    "TOTAL:               1.00  (100%)"
)

body("Metric Formulas", bold=True, color=DARK_BLUE, size=12)
code_block(
    "schedule_completeness  = assigned_count / total_flights\n\n"
    "conflict_free_ratio    = 1.0 − (conflict_count / total_flights)\n\n"
    "priority_handling      = priority_on_time_count / priority_total\n\n"
    "delay_efficiency       = 1.0 − (total_delay / delay_budget)\n\n"
    "fuel_efficiency        = 1.0 − (total_fuel_burn / fuel_budget)\n\n"
    "fairness               = 1.0 − (std_dev(per_airline_avg_delay) / tolerance)\n\n"
    "normalized_score       = clamp(Σ weight_i × metric_i, 0.01, 0.99)"
)

body("Composite Grader Formula", bold=True, color=DARK_BLUE, size=12)
code_block(
    "final_score = clamp(\n"
    "    min(safety_gate_ceiling,\n"
    "        0.30 × priority_score  +  0.70 × efficiency_score),\n"
    "    0.01, 0.99\n"
    ")"
)

body("Capacity Gap Formula", bold=True, color=DARK_BLUE, size=12)
code_block(
    "base_gap = max(2, round(60 / hourly_capacity))\n"
    "gap      = round(base_gap × weather_penalty_factor)"
)

body("Potential-Based Reward", bold=True, color=DARK_BLUE, size=12)
code_block(
    "r_t = score_t − score_{t-1}          # step reward\n"
    "Σ r_t = score_final − score_initial  # episode return\n"
    "      = score_final  (since score_initial = 0)"
)

body("Wake Separation Numbers (ICAO-based)", bold=True, color=DARK_BLUE, size=12)
code_block(
    "H→H=4, H→M=5, H→L=6\n"
    "M→H=3, M→M=3, M→L=4\n"
    "L→H=3, L→M=3, L→L=3   (all values in minutes)"
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 12 — CODING PATTERNS
# ═══════════════════════════════════════════════════════════════════════════════
heading("12. Coding Patterns You Must Know", level=1)

heading("12.1 — Minimal OpenEnv Environment (copy-paste template)", level=2, color=DARK_BLUE)
code_block(
    "from openenv import Environment\n"
    "from pydantic import BaseModel\n"
    "from typing import Optional\n\n"
    "class MyAction(BaseModel):\n"
    "    decisions: list[dict]\n"
    "    rationale: str\n"
    "    commit: bool = False\n\n"
    "class MyObservation(BaseModel):\n"
    "    task_description: str\n"
    "    current_score: float\n"
    "    reward: float\n"
    "    done: bool\n"
    "    diagnostics: list[str] = []\n\n"
    "class MyState(BaseModel):\n"
    "    episode_id: str\n"
    "    step_count: int\n"
    "    best_score: float\n\n"
    "class MyEnvironment(Environment[MyAction, MyObservation, MyState]):\n"
    "    SUPPORTS_CONCURRENT_SESSIONS = True\n\n"
    "    def reset(self, task_id=None, **kwargs) -> MyObservation:\n"
    "        self._score = 0.0\n"
    "        self._steps = 0\n"
    "        return MyObservation(task_description='...', current_score=0.0,\n"
    "                             reward=0.0, done=False)\n\n"
    "    def step(self, action: MyAction) -> MyObservation:\n"
    "        new_score   = self._simulate(action)\n"
    "        reward      = new_score - self._score\n"
    "        self._score = new_score\n"
    "        self._steps += 1\n"
    "        done = action.commit or self._steps >= MAX_STEPS\n"
    "        return MyObservation(task_description='...', current_score=new_score,\n"
    "                             reward=reward, done=done)\n\n"
    "    @property\n"
    "    def state(self) -> MyState:\n"
    "        return MyState(episode_id='ep1', step_count=self._steps, best_score=self._score)"
)

heading("12.2 — Pydantic Model Pattern", level=2, color=DARK_BLUE)
code_block(
    "from pydantic import BaseModel, Field\n"
    "from enum import Enum\n\n"
    "class PriorityClass(str, Enum):\n"
    "    EMERGENCY  = 'emergency'\n"
    "    NORMAL     = 'normal'\n\n"
    "class FlightRecord(BaseModel):\n"
    "    flight_id:  str\n"
    "    priority:   PriorityClass\n"
    "    earliest:   int                     # minute\n"
    "    latest:     int\n"
    "    passengers: int = 0\n"
    "    metadata:   dict = Field(default_factory=dict)"
)

heading("12.3 — 3-Layer Grader Template", level=2, color=DARK_BLUE)
code_block(
    "def grade(metrics: dict) -> float:\n"
    "    # Layer 1: Safety gate\n"
    "    ceiling = 1.0\n"
    "    if metrics['conflicts'] > 0:\n"
    "        ceiling = max(0.10, 0.40 - 0.05 * (metrics['conflicts'] - 1))\n\n"
    "    # Layer 2: Priority score (30%)\n"
    "    priority = metrics.get('priority_handling', 0.0)\n\n"
    "    # Layer 3: Efficiency score (70%)\n"
    "    efficiency = (\n"
    "        0.35 * metrics.get('delay_efficiency', 0.0)\n"
    "      + 0.25 * metrics.get('fuel_efficiency',  0.0)\n"
    "      + 0.20 * metrics.get('fairness',         0.0)\n"
    "      + 0.20 * metrics.get('connection_impact',0.0)\n"
    "    )\n\n"
    "    raw = 0.30 * priority + 0.70 * efficiency\n"
    "    return max(0.01, min(0.99, min(ceiling, raw)))"
)

heading("12.4 — FastAPI App Wiring", level=2, color=DARK_BLUE)
code_block(
    "from openenv import create_app\n"
    "from .my_environment import MyEnvironment\n\n"
    "app = create_app(\n"
    "    MyEnvironment,\n"
    "    title='My OpenEnv',\n"
    "    max_environments=8,\n"
    ")\n\n"
    "# openenv core auto-adds /reset, /step, /state, /health"
)

heading("12.5 — Inference Loop Template", level=2, color=DARK_BLUE)
code_block(
    "import asyncio, httpx\n\n"
    "async def run_inference(base_url: str, task_id: str):\n"
    "    async with httpx.AsyncClient() as client:\n"
    "        # Reset\n"
    "        obs = (await client.post(f'{base_url}/reset',\n"
    "                                  json={'task_id': task_id})).json()\n"
    "        print(f'[START] task={task_id}')\n\n"
    "        for step in range(MAX_STEPS):\n"
    "            action = get_action(obs)            # your agent\n"
    "            result = (await client.post(f'{base_url}/step',\n"
    "                                         json=action)).json()\n"
    "            reward = result['reward']\n"
    "            done   = result['done']\n"
    "            print(f'[STEP] step={step} reward={reward:.4f} done={done}')\n"
    "            obs = result\n"
    "            if done:\n"
    "                break\n\n"
    "        print(f'[END] task={task_id} score={obs[\"current_score\"]:.4f}')"
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  FINAL SECTION — ROUND 2 CHECKLIST
# ═══════════════════════════════════════════════════════════════════════════════
heading("Round 2 Action Checklist", level=1)
body("Do these IN ORDER when you get the new problem statement:", bold=True, color=ACCENT_ORG, size=13)

steps = [
    "READ the problem statement carefully. Identify: (a) the domain objects, (b) the hard constraints, (c) the soft optimisation goals.",
    "DESIGN the constraint hierarchy: what must NEVER be violated (safety gate), what is critical (priority, 30%), what is efficiency (70%).",
    "DEFINE models.py first: Action, Observation, State Pydantic models. This forces you to think about what the agent sees and does.",
    "WRITE constants.py: domain-specific constants, score weights, constraint thresholds.",
    "WRITE engine.py: deterministic simulate() function that takes an action and returns metrics.",
    "WRITE graders.py: 3-layer gated grader using the template from Section 12.3.",
    "WRITE tasks.py: at least 3 tasks (easy, medium, hard). Write render_task_briefing() for each.",
    "WRITE planner.py: greedy heuristic baseline. Sort by priority, assign greedily respecting constraints.",
    "WRITE server/atc_environment.py → rename: implement reset() and step() using potential-based reward.",
    "UPDATE openenv.yaml: new task IDs, action/observation space references.",
    "TEST: run pytest; check score bounds (0.01, 0.99), determinism, baseline > random.",
    "DEPLOY: docker build + HF Space push (or local uvicorn for demo).",
]
for i, step in enumerate(steps, 1):
    numbered(step)

doc.add_paragraph()
callout("REMEMBER:", "Your biggest advantage is that you have already done this once. "
        "Most Round 2 teams will spend 50% of their time on infrastructure. "
        "You can spend 100% on domain logic because your infrastructure is battle-tested.", "E8F8E8")

# ─── Save ──────────────────────────────────────────────────────────────────────
output_path = r"g:\ATC\OpenEnv_Hackathon_Complete_Guide.docx"
doc.save(output_path)
print(f"Document saved to: {output_path}")
